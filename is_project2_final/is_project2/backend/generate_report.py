"""
Script to generate a comprehensive Word document report for the Eunoia Chatbot
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading(doc, text, level=1):
    """Add a heading with custom formatting"""
    heading = doc.add_heading(text, level=level)
    heading.style.font.size = Pt(14 if level == 1 else 12)
    return heading

def add_table_from_data(doc, headers, rows, title=None):
    """Add a formatted table to the document"""
    if title:
        doc.add_paragraph(title, style='Heading 3')
    
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Add headers
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
    
    # Add data rows
    for row_idx, row_data in enumerate(rows, start=1):
        row_cells = table.rows[row_idx].cells
        for col_idx, cell_data in enumerate(row_data):
            row_cells[col_idx].text = str(cell_data)
    
    doc.add_paragraph()  # Add spacing after table

def create_report():
    """Create the comprehensive Word document report"""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading('Eunoia Chatbot: Comprehensive Technical Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph('A Detailed Analysis of Architecture, Workflow, Metrics, and Design Decisions')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.italic = True
    subtitle.runs[0].font.size = Pt(12)
    
    doc.add_paragraph()  # Spacing
    
    # Executive Summary
    add_heading(doc, 'Executive Summary', 1)
    doc.add_paragraph(
        'Eunoia is a bilingual (English/Kiswahili) menstrual health chatbot that uses RAG (Retrieval-Augmented Generation) '
        'with a fine-tuned Flan-T5 model. The system supports both English and Kiswahili languages, detects user emotions, '
        'and provides empathetic, medically safe responses tailored to the Kenyan context.'
    )
    
    doc.add_page_break()
    
    # Section 1: Complete User Journey
    add_heading(doc, '1. Complete User Journey: Question to Response', 1)
    
    add_heading(doc, '1.1 Frontend Flow (React)', 2)
    doc.add_paragraph('1. User Input: User types question in ChatPage.jsx')
    doc.add_paragraph('2. Language Detection: Frontend sends language preference (from LanguageContext)')
    doc.add_paragraph('3. API Call: POST request to http://127.0.0.1:5000/chat with:')
    p = doc.add_paragraph('   • message: user input text', style='List Bullet')
    p = doc.add_paragraph('   • user_id: Firebase UID', style='List Bullet')
    p = doc.add_paragraph('   • conversation_id: current conversation ID', style='List Bullet')
    p = doc.add_paragraph('   • language: user preference ("en" or "sw")', style='List Bullet')
    
    add_heading(doc, '1.2 Backend Processing Pipeline (fallback1.py)', 2)
    
    steps = [
        ('Step 1: Input Validation & Conversation Management', 
         'Validates input, creates conversation if needed, loads conversation history (last 5 messages)'),
        ('Step 2: Language Detection', 
         'Uses langdetect library to detect input language. Respects user preference: if preference is "sw", always respond in Swahili. If user writes in Swahili but preference is English, responds in English.'),
        ('Step 3: Special Case Handling', 
         'Handles greetings, off-topic questions, and bot identity questions with predefined responses.'),
        ('Step 4: Emotion Detection', 
         'Detects emotion (pain, anxious, sad, neutral) using keyword matching in both English and Swahili.'),
        ('Step 5: Query Translation (Swahili Mode)', 
         'If language == "sw": Translates Swahili query → English using Bildad/Swahili-English_Translation model. Includes preprocessing for common patterns and post-processing to fix translation errors.'),
        ('Step 6: Context Retrieval (RAG)', 
         'Uses FAISS vector search with SentenceTransformer embeddings. Retrieves top 5 relevant contexts from knowledge base. Filters irrelevant content and Indian program references.'),
        ('Step 7: Context Summarization', 
         'Truncates context to 120 words to fit within token limits while keeping most relevant information.'),
        ('Step 8: Response Generation', 
         'Uses fine-tuned Flan-T5 model with carefully crafted prompts including system instructions, emotion guidance, context, and question. Generation parameters: repetition_penalty=1.8, num_beams=4, temperature=0.85.'),
        ('Step 9: Response Translation (Swahili Mode)', 
         'If language == "sw": Translates English response → Swahili using Helsinki-NLP/opus-mt-en-sw. Applies naturalization to make Swahili more casual/conversational (Kenyan Kiswahili style).'),
        ('Step 10: Response Validation & Cleaning', 
         'Validates response for: instruction echoes, medical safety, duplicates, contradictions, typos, generic closings, minimum length, warm tone, irrelevant context removal.'),
        ('Step 11: Fallback System', 
         'If response fails quality checks, uses empathetic fallback module that structures response as: Validation → Explanation → Tips → Closing.'),
        ('Step 12: Save & Return', 
         'Saves to conversation history and returns JSON response with response text, detected emotion, language, and conversation_id.')
    ]
    
    for i, (title, desc) in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {title}', style='Heading 3')
        doc.add_paragraph(desc)
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # Section 2: Technologies & Rationale
    add_heading(doc, '2. Technologies & Rationale', 1)
    
    technologies = [
        ('Core Model: Flan-T5 (Fine-tuned)', 
         'Why: Excellent instruction-following capabilities, efficient, fine-tuned on menstrual health data. Size: Base model, suitable for CPU/GPU deployment. Fine-tuning: Trained on menstrual_data.csv Q&A pairs.'),
        ('Embeddings: SentenceTransformer ("all-MiniLM-L6-v2")', 
         'Why: Fast, 384-dimensional embeddings, good semantic search quality. Alternative: Larger models (slower) or domain-specific (requires training).'),
        ('Vector Database: FAISS (Facebook AI Similarity Search)', 
         'Why: Fast similarity search, supports L2 distance, efficient indexing. Index type: IndexFlatL2 (exact search, good for this corpus size). Alternative: ChromaDB, Pinecone (overkill for this scale).'),
        ('Knowledge Base: menstrual_data.csv', 
         'Structure: Question-Answer pairs. Size: ~22,305 pairs (test set: 4,461). Swahili version: menstrual_data_sw.csv (pre-translated).'),
        ('Translation: English → Swahili (Helsinki-NLP/opus-mt-en-sw)', 
         'Why: MarianMT architecture, good quality, widely used. Naturalization: Custom post-processing for Kenyan Kiswahili style.'),
        ('Translation: Swahili → English (Bildad/Swahili-English_Translation)', 
         'Why: Better for Swahili → English than reverse model. Preprocessing: Direct mappings for common queries. Post-processing: Fixes common errors.'),
        ('Language Detection: langdetect', 
         'Why: Lightweight, supports Swahili. Alternative: polyglot (heavier), fasttext (requires model).'),
        ('Backend Framework: Flask', 
         'Why: Simple, Python-friendly, easy to deploy. Alternative: FastAPI (more modern, async), Django (overkill).'),
        ('Frontend: React + Vite', 
         'Why: Modern, component-based, good UX. State: React Context for language/dark mode. Routing: React Router.'),
        ('Authentication: Firebase Auth', 
         'Why: Managed authentication, secure, easy integration. User management: Per-user conversation files.')
    ]
    
    for title, desc in technologies:
        doc.add_paragraph(title, style='Heading 3')
        doc.add_paragraph(desc)
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # Section 3: Evaluation Metrics
    add_heading(doc, '3. Evaluation Metrics & Results', 1)
    
    add_heading(doc, '3.1 Retrieval Metrics (RAG)', 2)
    
    retrieval_headers = ['Metric', 'Top-1', 'Top-3', 'Top-5', 'Top-10']
    retrieval_rows = [
        ['Precision@K', '0.052', '0.182', '0.318', '0.526'],
        ['Recall@K', '0.052', '0.182', '0.318', '0.526'],
        ['MRR@K', '0.052', '0.106', '0.136', '0.164'],
        ['Context Relevance', '0.626', '0.632', '0.634', '0.635']
    ]
    add_table_from_data(doc, retrieval_headers, retrieval_rows)
    
    doc.add_paragraph('Interpretation:', style='Heading 4')
    doc.add_paragraph('• Precision@5: 31.8% of queries find the exact answer in top 5 results')
    doc.add_paragraph('• Context Relevance: ~0.63 semantic similarity (good quality)')
    doc.add_paragraph('• Top-10: 52.6% precision (reasonable for diverse questions)')
    doc.add_paragraph()
    
    add_heading(doc, '3.2 Generation Metrics (With RAG vs Without RAG)', 2)
    
    generation_headers = ['Metric', 'With RAG', 'Without RAG', 'Improvement']
    generation_rows = [
        ['BLEU', '0.328', '0.019', '+1,625%'],
        ['ROUGE-1', '0.474', '0.175', '+171%'],
        ['ROUGE-2', '0.390', '0.046', '+745%'],
        ['ROUGE-L', '0.442', '0.139', '+217%'],
        ['METEOR', '0.508', '0.152', '+234%'],
        ['Semantic Similarity', '0.778', '0.551', '+41%']
    ]
    add_table_from_data(doc, generation_headers, generation_rows)
    
    doc.add_paragraph('Key Findings:', style='Heading 4')
    doc.add_paragraph('• RAG dramatically improves all metrics')
    doc.add_paragraph('• BLEU: 0.328 (moderate, but much better than baseline)')
    doc.add_paragraph('• ROUGE-L: 0.442 (good overlap with reference)')
    doc.add_paragraph('• METEOR: 0.508 (good semantic match)')
    doc.add_paragraph('• Semantic Similarity: 0.778 (strong semantic alignment)')
    doc.add_paragraph()
    
    add_heading(doc, '3.3 Model Performance Analysis', 2)
    doc.add_paragraph('Strengths:', style='Heading 4')
    doc.add_paragraph('• High semantic similarity (0.778)')
    doc.add_paragraph('• Good ROUGE scores (0.39-0.47)')
    doc.add_paragraph('• RAG provides substantial improvement')
    doc.add_paragraph()
    doc.add_paragraph('Challenges:', style='Heading 4')
    doc.add_paragraph('• BLEU variability (std: 0.373) - some responses very different from references')
    doc.add_paragraph('• Repetition issues (addressed with repetition_penalty=1.8)')
    doc.add_paragraph('• Instruction echoing (mitigated with validation)')
    
    doc.add_page_break()
    
    # Section 4: Architecture Decisions
    add_heading(doc, '4. Architecture Decisions & Rationale', 1)
    
    decisions = [
        ('Bilingual Architecture', 
         'Approach: Generate in English, translate to Swahili. Why: English corpus is larger and more reliable, better model performance in English, translation quality is more consistent. Alternative: Train separate Swahili model (requires more data/resources).'),
        ('Multi-Conversation System', 
         'Design: Per-user conversation files with conversation IDs. Why: Privacy (each user has isolated data), scalability (easy to migrate to database later), simplicity (JSON files are easy to debug). Structure: JSON file per user with nested conversations.'),
        ('Context Summarization', 
         'Why: Limits prompt length, keeps most relevant info. Max words: 120. Prevents token limit issues, focuses on key information.'),
        ('Emotion-Aware Responses', 
         'Why: Menstrual health is sensitive topic; empathy matters. Emotion detection: Keyword-based (fast, reliable). Emotion-specific openings/closings. Tone adaptation.'),
        ('Medical Safety Validation', 
         'Why: Critical for health information. Unsafe pattern detection. Tampon safety corrections. Contradiction removal. Fallback to safe responses.'),
        ('Fallback System', 
         'Why: Ensures quality responses even if generation fails. Empathetic fallback module. Structure: Validation → Explanation → Tips → Closing. Always provides helpful response.')
    ]
    
    for title, desc in decisions:
        doc.add_paragraph(title, style='Heading 3')
        doc.add_paragraph(desc)
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # Section 5: Limitations & Improvements
    add_heading(doc, '5. System Limitations & Future Improvements', 1)
    
    add_heading(doc, 'Current Limitations:', 2)
    doc.add_paragraph('1. Retrieval precision: 31.8% at top-5 (could be higher)')
    doc.add_paragraph('2. BLEU variability: Some responses differ significantly from references')
    doc.add_paragraph('3. Translation quality: Depends on model quality (some errors)')
    doc.add_paragraph('4. Emotion detection: Keyword-based (could use ML model)')
    doc.add_paragraph('5. Context filtering: May remove relevant information')
    doc.add_paragraph()
    
    add_heading(doc, 'Potential Improvements:', 2)
    doc.add_paragraph('1. Better embeddings: Domain-specific model or larger general model')
    doc.add_paragraph('2. Re-ranking: Use cross-encoder to re-rank retrieved contexts')
    doc.add_paragraph('3. Hybrid search: Combine semantic + keyword search')
    doc.add_paragraph('4. Emotion model: Train/use ML-based emotion detection')
    doc.add_paragraph('5. Translation fine-tuning: Fine-tune on menstrual health domain')
    doc.add_paragraph('6. Database migration: Move from JSON files to PostgreSQL/MongoDB')
    doc.add_paragraph('7. Caching: Cache common queries/translations')
    doc.add_paragraph('8. A/B testing: Test different prompt strategies')
    
    doc.add_page_break()
    
    # Section 6: Deployment
    add_heading(doc, '6. Deployment Architecture', 1)
    
    add_heading(doc, 'Current Setup:', 2)
    doc.add_paragraph('• Backend: Flask on localhost:5000')
    doc.add_paragraph('• Frontend: React/Vite dev server')
    doc.add_paragraph('• Storage: Local JSON files')
    doc.add_paragraph('• Models: Local (downloaded on first run)')
    doc.add_paragraph()
    
    add_heading(doc, 'Production Considerations:', 2)
    doc.add_paragraph('1. Model serving: Use GPU server or cloud (AWS/GCP)')
    doc.add_paragraph('2. API: Deploy Flask with Gunicorn/uWSGI')
    doc.add_paragraph('3. Frontend: Build and serve via Nginx/CDN')
    doc.add_paragraph('4. Database: Migrate to PostgreSQL for conversations')
    doc.add_paragraph('5. Caching: Redis for common queries')
    doc.add_paragraph('6. Monitoring: Logging, error tracking (Sentry)')
    doc.add_paragraph('7. Scaling: Load balancer, multiple workers')
    
    doc.add_page_break()
    
    # Section 7: Conclusion
    add_heading(doc, '7. Conclusion', 1)
    doc.add_paragraph(
        'Eunoia successfully implements a sophisticated RAG-based chatbot system with bilingual support, emotion detection, '
        'and medical safety validation. The system demonstrates significant improvements with RAG (BLEU +1,625%, ROUGE-L +217%). '
        'The architecture prioritizes medical safety, empathy, and bilingual accessibility for Kenyan users.'
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'The system balances performance, safety, and user experience, with clear paths for future improvements in retrieval, '
        'translation, and emotion detection.'
    )
    
    # Save document
    output_path = 'Eunoia_Chatbot_Technical_Report.docx'
    doc.save(output_path)
    print(f'✅ Report generated successfully: {output_path}')
    return output_path

if __name__ == '__main__':
    try:
        create_report()
    except ImportError:
        print('❌ Error: python-docx library not installed')
        print('   Please install it with: pip install python-docx')
    except Exception as e:
        print(f'❌ Error generating report: {e}')

